"""
Export and refinement engine for NTRO Platform.
Handles PDF/DOCX generation and content refinement via LLM.
"""

import io
import json
import logging
import os
from datetime import datetime
from typing import Any

from openai import AsyncOpenAI
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from prompts import IngestionParameters

logger = logging.getLogger(__name__)


class PDFExporter:
    """Export deliverables to PDF format."""

    @staticmethod
    def generate_pdf(
        deliverables: dict[str, Any],
        parameters: IngestionParameters,
    ) -> bytes:
        """
        Generate PDF from deliverables.

        Args:
            deliverables: Dictionary of format -> content
            parameters: Generation parameters for metadata

        Returns:
            PDF file as bytes
        """
        try:
            pdf_buffer = io.BytesIO()

            # Create PDF document
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=0.75 * inch,
                leftMargin=0.75 * inch,
                topMargin=1 * inch,
                bottomMargin=0.75 * inch,
            )

            story = []

            # Title page
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=24,
                textColor="rgb(25, 118, 210)",
                spaceAfter=30,
                alignment=1,  # Center
            )

            story.append(Spacer(1, 1.5 * inch))
            story.append(
                Paragraph("NTRO Platform - Deliverables Report", title_style)
            )
            story.append(
                Paragraph(
                    f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                    styles["Normal"],
                )
            )
            story.append(
                Paragraph(
                    f"Audience: {parameters.target_audience} | Tone: {parameters.tone}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 0.5 * inch))

            # Content sections
            for format_name, content in deliverables.items():
                story.append(PageBreak())
                story.append(
                    Paragraph(format_name, styles["Heading2"])
                )
                story.append(Spacer(1, 0.2 * inch))

                # Handle different content types
                if isinstance(content, list):
                    # Twitter/X thread
                    for i, tweet in enumerate(content, 1):
                        story.append(
                            Paragraph(
                                f"<b>Tweet {i}:</b> {tweet}",
                                styles["BodyText"],
                            )
                        )
                        story.append(Spacer(1, 0.15 * inch))
                else:
                    # Regular text content
                    story.append(Paragraph(str(content), styles["BodyText"]))

                story.append(Spacer(1, 0.3 * inch))

            # Build PDF
            doc.build(story)
            pdf_buffer.seek(0)
            return pdf_buffer.getvalue()

        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}", exc_info=True)
            raise

    @staticmethod
    def get_mime_type() -> str:
        """Get MIME type for PDF."""
        return "application/pdf"


class DOCXExporter:
    """Export deliverables to DOCX (Word) format."""

    @staticmethod
    def generate_docx(
        deliverables: dict[str, Any],
        parameters: IngestionParameters,
    ) -> bytes:
        """
        Generate DOCX from deliverables.

        Args:
            deliverables: Dictionary of format -> content
            parameters: Generation parameters for metadata

        Returns:
            DOCX file as bytes
        """
        try:
            doc = Document()

            # Add title
            title = doc.add_heading("NTRO Platform - Deliverables Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata
            meta = doc.add_paragraph()
            meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
            meta.add_run(
                f"Audience: {parameters.target_audience} | Tone: {parameters.tone}"
            )

            doc.add_paragraph()  # Spacing

            # Add deliverables
            for format_name, content in deliverables.items():
                doc.add_heading(format_name, level=1)

                if isinstance(content, list):
                    # Twitter/X thread
                    for i, tweet in enumerate(content, 1):
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(f"Tweet {i}: {tweet}")
                else:
                    # Regular content
                    doc.add_paragraph(str(content))

                doc.add_paragraph()  # Spacing

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            return docx_buffer.getvalue()

        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}", exc_info=True)
            raise

    @staticmethod
    def get_mime_type() -> str:
        """Get MIME type for DOCX."""
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class ContentRefiner:
    """Refine generated content using LLM."""

    @staticmethod
    async def refine_content(
        original_content: str,
        instruction: str,
        format_type: str,
        parameters: IngestionParameters,
    ) -> dict[str, Any]:
        """
        Refine content based on user instruction.

        Args:
            original_content: Original generated content
            instruction: Refinement instruction (e.g., "Make more concise")
            format_type: Type of deliverable
            parameters: Generation parameters

        Returns:
            Dictionary with refined_content and change_summary
        """
        try:
            api_key = os.getenv("GROQ_API_KEY")
            if not api_key:
                raise ValueError("GROQ_API_KEY not configured")

            client = AsyncOpenAI(
                api_key=api_key,
                base_url="https://api.groq.com/openai/v1"
            )

            refinement_prompt = f"""You are a content refinement specialist. 
            
Current content ({format_type}):
{original_content}

Refinement instruction: {instruction}

Refine the content according to the instruction while maintaining the core message and format requirements.
Return ONLY the refined content without any explanations."""

            response = await client.chat.completions.create(
                model="openai/gpt-oss-120b",
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional content editor specializing in {format_type} format."
                    },
                    {"role": "user", "content": refinement_prompt},
                ],
                temperature=0.5,
                max_tokens=2000,
            )

            refined_content = response.choices[0].message.content or ""

            # Generate change summary
            change_summary = f"Applied: {instruction}"

            logger.info(f"Content refined successfully for {format_type}")

            return {
                "status": "success",
                "original_length": len(original_content),
                "refined_length": len(refined_content),
                "refined_content": refined_content,
                "change_summary": change_summary,
            }

        except Exception as e:
            logger.error(f"Content refinement failed: {str(e)}", exc_info=True)
            return {
                "status": "error",
                "error": f"Refinement failed: {str(e)}",
            }
