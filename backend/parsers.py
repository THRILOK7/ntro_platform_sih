"""
Robust modular document parsing utility for ingesting multiple file formats.
Supports PDF (with OCR fallback), DOCX, images (with OCR), and audio/video transcription.
"""

import io
import os
import tempfile
from pathlib import Path
from typing import Optional
import logging

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from openai import OpenAI

logger = logging.getLogger(__name__)

# Windows Tesseract auto-fallback configuration
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def _ensure_tesseract_available() -> None:
    """
    Check if pytesseract can locate tesseract.
    If not, explicitly assign the Windows default path.
    """
    try:
        pytesseract.get_tesseract_version()
    except pytesseract.TesseractNotFoundError:
        if Path(TESSERACT_PATH).exists():
            pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
            logger.info(f"Tesseract configured at: {TESSERACT_PATH}")
        else:
            logger.warning(
                f"Tesseract not found at {TESSERACT_PATH}. "
                "OCR functionality may be unavailable."
            )


def parse_pdf(file_bytes: bytes) -> str:
    """
    Parse PDF content using PyMuPDF (fitz).
    If extracted text is < 50 chars, apply OCR to page images.

    Args:
        file_bytes: Raw PDF file bytes

    Returns:
        Extracted text from PDF

    Raises:
        ValueError: If PDF cannot be opened or processed
    """
    try:
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        extracted_text = ""

        for page_num, page in enumerate(pdf_doc):
            text = page.get_text()
            extracted_text += text

        pdf_doc.close()

        # If text extraction yielded minimal content, apply OCR
        if len(extracted_text.strip()) < 50:
            logger.warning(
                f"PDF text extraction yielded < 50 chars. Attempting OCR fallback."
            )
            _ensure_tesseract_available()
            extracted_text = _ocr_pdf_pages(file_bytes)

        return extracted_text.strip()

    except Exception as e:
        raise ValueError(f"PDF parsing failed: {str(e)}") from e


def _ocr_pdf_pages(file_bytes: bytes) -> str:
    """
    Apply OCR to all pages of a PDF document.

    Args:
        file_bytes: Raw PDF file bytes

    Returns:
        OCR'd text from all pages
    """
    try:
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        ocr_text = ""

        for page_num, page in enumerate(pdf_doc):
            # Render page to image with high DPI for better OCR
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_text = pytesseract.image_to_string(img)
            ocr_text += f"\n--- Page {page_num + 1} ---\n{page_text}"

        pdf_doc.close()
        return ocr_text.strip()

    except Exception as e:
        logger.error(f"PDF OCR failed: {str(e)}")
        raise ValueError(f"PDF OCR processing failed: {str(e)}") from e


def parse_docx(file_bytes: bytes) -> str:
    """
    Parse DOCX content using python-docx.
    Extracts paragraph text and table cells.

    Args:
        file_bytes: Raw DOCX file bytes

    Returns:
        Extracted text from DOCX document

    Raises:
        ValueError: If DOCX cannot be opened or processed
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        extracted_text = ""

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text += para.text + "\n"

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    extracted_text += " | ".join(row_cells) + "\n"

        return extracted_text.strip()

    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {str(e)}") from e


def parse_image(file_bytes: bytes) -> str:
    """
    Parse image content using Tesseract OCR.

    Args:
        file_bytes: Raw image file bytes

    Returns:
        OCR'd text from image

    Raises:
        ValueError: If image cannot be opened or OCR fails
    """
    try:
        _ensure_tesseract_available()
        img = Image.open(io.BytesIO(file_bytes))
        extracted_text = pytesseract.image_to_string(img)
        return extracted_text.strip()

    except Exception as e:
        raise ValueError(f"Image parsing failed: {str(e)}") from e


def parse_audio_video(file_bytes: bytes, filename: str) -> str:
    """
    Transcribe audio/video using OpenAI Whisper API.
    Writes bytes to a temporary file, processes with Whisper, and cleans up.

    Args:
        file_bytes: Raw audio/video file bytes
        filename: Original filename (used for temp file suffix)

    Returns:
        Transcribed text from audio/video

    Raises:
        ValueError: If transcription fails

    MANDATORY: try...finally block ensures temp file cleanup on all code paths.
    """
    temp_file_path = None

    try:
        # Determine file extension
        ext = Path(filename).suffix or ".mp3"

        # Create secure temporary file
        with tempfile.NamedTemporaryFile(
            suffix=ext, delete=False, prefix="ntro_audio_"
        ) as temp_file:
            temp_file.write(file_bytes)
            temp_file_path = temp_file.name

        # Transcribe using OpenAI Whisper
        with open(temp_file_path, "rb") as audio_file:
            client = OpenAI()
            transcript = client.audio.transcriptions.create(
                model="whisper-1", file=audio_file
            )

        return transcript.text.strip()

    except Exception as e:
        raise ValueError(f"Audio/video transcription failed: {str(e)}") from e

    finally:
        # MANDATORY: Ensure temp file is always deleted
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.debug(f"Cleaned up temporary file: {temp_file_path}")
            except Exception as cleanup_error:
                logger.error(
                    f"Failed to clean up temp file {temp_file_path}: {cleanup_error}"
                )


def extract_content(filename: str, file_bytes: bytes) -> str:
    """
    Route file content extraction to appropriate parser based on file extension.

    Args:
        filename: Original filename (with extension)
        file_bytes: Raw file bytes

    Returns:
        Extracted content as string

    Raises:
        ValueError: If file type is unsupported or parsing fails
    """
    file_ext = Path(filename).suffix.lower()

    # PDF
    if file_ext == ".pdf":
        return parse_pdf(file_bytes)

    # DOCX
    elif file_ext == ".docx":
        return parse_docx(file_bytes)

    # Images
    elif file_ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff"}:
        return parse_image(file_bytes)

    # Audio/Video
    elif file_ext in {
        ".mp3",
        ".wav",
        ".m4a",
        ".flac",
        ".ogg",
        ".mp4",
        ".webm",
        ".mov",
        ".avi",
    }:
        return parse_audio_video(file_bytes, filename)

    # Unsupported
    else:
        raise ValueError(
            f"Unsupported file type: {file_ext}. "
            "Supported: PDF, DOCX, PNG, JPG, JPEG, GIF, BMP, WEBP, TIFF, "
            "MP3, WAV, M4A, FLAC, OGG, MP4, WEBM, MOV, AVI"
        )
